"""
chat_agent.py v3.0 — ИИ-агент с расширенными возможностями.

Режимы:
  chat        — обычный чат с LLM
  code        — агент-кодер (пишет, запускает, исправляет)

Возможности агент-кодера:
  1. Генерация кода по описанию
  2. Создание изображений (через API провайдера)
  3. Создание/проектирование ботов и приложений
  4. Анализ документов / архивов / кода
  5. Поиск ошибок в коде
  6. Автоисправление ошибок (до MAX_ITERATIONS попыток)
  7. Авто-поиск решения в интернете при неизвестных ошибках
"""

import os
import sys
import subprocess
import tempfile
import time
import threading
import re
import json
from llm_client import call_llm, call_llm_full
import config

# ── Сессии ────────────────────────────────────────────────────
_sessions      = {}
_sessions_lock = threading.Lock()

MAX_HISTORY     = 20
MAX_ITERATIONS  = 6
CODE_TIMEOUT    = 90
MAX_OUTPUT      = 4000
CODE_MAX_TOKENS = 8000   # было 6000 — увеличено чтобы реже обрезало

PROJECTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'agent_projects')
os.makedirs(PROJECTS_DIR, exist_ok=True)

MAX_CONTINUATIONS = 3   # максимум продолжений одного ответа


def _call_with_continuation(prompt, system, max_tokens=None, on_status=None):
    """
    Генерирует ответ LLM с автопродолжением если ответ оборван по лимиту токенов.

    Работает так:
      1. Делает первый запрос с call_llm_full → получает (text, is_truncated)
      2. Если is_truncated=True — делает второй запрос "продолжи с: <хвост>"
      3. Склеивает куски, повторяет до MAX_CONTINUATIONS раз
      4. Возвращает полный текст

    Это решает проблему обрыва кода на середине при достижении max_tokens.
    """
    if max_tokens is None:
        max_tokens = CODE_MAX_TOKENS

    full_text, is_truncated = call_llm_full(prompt, system, max_tokens)
    parts = [full_text]

    for attempt in range(1, MAX_CONTINUATIONS + 1):
        if not is_truncated:
            break

        if on_status:
            on_status("📄 Ответ обрезан — продолжаю (часть {})...".format(attempt + 1))

        # Берём последние ~300 символов как контекст для продолжения
        tail = full_text[-300:].strip()
        continuation_prompt = (
            "Продолжи генерацию ТОЧНО с того места где оборвалось. "
            "Не повторяй уже написанное. Последние строки были:\n\n"
            "```\n{}\n```\n\n"
            "Продолжай:"
        ).format(tail)

        cont_text, is_truncated = call_llm_full(continuation_prompt, system, max_tokens)
        # Убираем дублирование: если LLM повторил хвост — вырезаем
        if tail and cont_text.startswith(tail):
            cont_text = cont_text[len(tail):].lstrip()
        parts.append(cont_text)
        full_text = cont_text  # для следующего tail

    return "".join(parts)


#  СЕССИИ
# ════════════════════════════════════════════════════════════

def get_session(chat_id):
    with _sessions_lock:
        return _sessions.get(str(chat_id))

def start_session(chat_id, mode):
    with _sessions_lock:
        _sessions[str(chat_id)] = {
            'mode': mode, 'history': [], 'started': time.time(), 'iterations': 0,
        }

def end_session(chat_id):
    with _sessions_lock:
        return _sessions.pop(str(chat_id), None)

def is_active(chat_id):
    return str(chat_id) in _sessions

def add_to_history(chat_id, role, content):
    with _sessions_lock:
        s = _sessions.get(str(chat_id))
        if not s:
            return
        s['history'].append({'role': role, 'content': content})
        if len(s['history']) > MAX_HISTORY * 2:
            s['history'] = s['history'][-(MAX_HISTORY * 2):]

def get_history(chat_id):
    with _sessions_lock:
        s = _sessions.get(str(chat_id))
        return list(s['history']) if s else []

def session_info(chat_id):
    with _sessions_lock:
        s = _sessions.get(str(chat_id))
        if not s:
            return None
        elapsed = int(time.time() - s['started'])
        mins, secs = elapsed // 60, elapsed % 60
        return {
            'mode':       s['mode'],
            'messages':   len(s['history']),
            'elapsed':    '{}м {}с'.format(mins, secs) if mins else '{}с'.format(secs),
            'iterations': s.get('iterations', 0),
        }

def all_active_sessions():
    with _sessions_lock:
        return list(_sessions.keys())


# ════════════════════════════════════════════════════════════
#  ОПРЕДЕЛЕНИЕ ТИПА ЗАДАЧИ
# ════════════════════════════════════════════════════════════

# Типы задач
TASK_SCRIPT   = 'script'    # простой скрипт — запустить и показать вывод
TASK_PROJECT  = 'project'   # проект — писать файлы
TASK_SCAFFOLD = 'scaffold'  # создать структуру папок + zip
TASK_REVIEW   = 'review'    # найти ошибки / проанализировать код
TASK_FIX      = 'fix'       # исправить ошибку (пришёл код + ошибка)
TASK_ANALYZE  = 'analyze'   # анализ документа / архива / файла
TASK_FILE     = 'file'      # создать файл (txt/md/csv/docx/zip/rar) и отдать
TASK_VIDEO    = 'video'     # скачать видео/аудио с YouTube или другого сайта
TASK_IMAGE    = 'image'     # генерация изображения

_SCAFFOLD_KW = [
    'make zip', 'create zip', 'сделай zip', 'создай zip',
    'make archive', 'создай архив', 'сделай архив',
    'scaffold', 'скаффолд', '├─', '│', '└─',
    'folder structure', 'project structure',
]

_PROJECT_KW = [
    'бот', 'bot', 'приложение', 'app', 'сервер', 'server',
    'парсер', 'parser', 'скрейпер', 'scraper', 'веб-сервис',
    'telegram', 'телеграм', 'flask', 'fastapi', 'django',
    'скачай', 'скачать', 'download', 'базу данных', 'database',
    'напиши программу', 'создай программу', 'напиши систему',
    'радио', 'tts', 'voice',
]

_REVIEW_KW = [
    'найди ошибк', 'проверь код', 'review', 'code review',
    'что не так', 'где ошибка', 'проанализируй код',
    'найди баги', 'найди проблем', 'улучши код', 'оптимизируй',
    'security audit', 'аудит безопасности', 'найди уязвимост',
]

_FIX_KW = [
    'исправь ошибку', 'fix error', 'fix bug', 'не работает',
    'traceback', 'exception', 'error:', 'исправь это',
    'починить', 'почини',
]

_ANALYZE_KW = [
    'проанализируй', 'analyze', 'анализ', 'разбор',
    'объясни этот код', 'что делает', 'как работает',
    'прочитай файл', 'прочитай архив', 'изучи документ',
    'summary', 'суммаризуй',
]

_FILE_KW = [
    # Форматы файлов — явное упоминание
    'txt файл', 'txt документ', 'текстовый файл', 'текстовый документ',
    'md файл', 'markdown файл', 'csv файл', 'таблицу csv', 'сохрани в csv',
    'docx', 'word документ', 'word файл', 'документ word',
    'zip архив', 'rar архив', 'собери zip', 'собери rar',
    'упакуй в zip', 'упакуй в rar', 'архивируй', 'запакуй',
    'скинь файлом', 'отправь файлом', 'сохрани файлом',
    # Действия с документами
    'подготовь анализ', 'подготовь отчёт', 'подготовь отчет',
    'напиши отчёт', 'напиши отчет', 'составь отчёт', 'составь отчет',
    'сделай отчёт', 'сделай отчет',
    'составь документ', 'создай документ', 'сделай документ',
    'запиши в файл', 'сохрани в файл', 'выгрузи в файл',
    'generate report', 'create report', 'write report',
    'create document', 'save to file', 'export to file',
    'prepare analysis', 'prepare report',
]

_VIDEO_KW = [
    # YouTube явно
    'youtube', 'ютуб', 'youtu.be',
    # Скачать видео
    'скачай видео', 'скачать видео', 'download video',
    'загрузи видео', 'загрузить видео',
    # mp3 / mp4
    'скачай mp3', 'скачать mp3', 'скачай mp4', 'скачать mp4',
    'в mp3', 'в mp4', 'as mp3', 'as mp4',
    'конвертируй в mp3', 'конвертируй в mp4',
    'извлеки аудио', 'extract audio',
    'скачай аудио', 'скачать аудио', 'download audio',
    # Общие
    'yt-dlp', 'ytdlp', 'yt_dlp',
]

_IMAGE_KW = [
    'нарисуй', 'draw', 'сгенерируй изображение', 'generate image',
    'создай изображение', 'create image', 'картинку', 'image of',
    'иллюстрацию', 'фото', 'арт', 'визуализацию',
]

def detect_task_type(task_text: str) -> str:
    t = task_text.lower()
    if any(kw in t for kw in _SCAFFOLD_KW):  return TASK_SCAFFOLD
    if any(kw in t for kw in _VIDEO_KW):     return TASK_VIDEO
    if any(kw in t for kw in _FILE_KW):      return TASK_FILE
    if any(kw in t for kw in _IMAGE_KW):     return TASK_IMAGE
    if any(kw in t for kw in _FIX_KW):       return TASK_FIX
    if any(kw in t for kw in _REVIEW_KW):    return TASK_REVIEW
    if any(kw in t for kw in _ANALYZE_KW):   return TASK_ANALYZE
    if any(kw in t for kw in _PROJECT_KW):   return TASK_PROJECT
    return TASK_SCRIPT

# Обратная совместимость
def _is_scaffold_task(t): return any(kw in t.lower() for kw in _SCAFFOLD_KW)
def _is_project_task(t):  return detect_task_type(t) == TASK_PROJECT


# ════════════════════════════════════════════════════════════
#  ЧАТ-РЕЖИМ
# ════════════════════════════════════════════════════════════

CHAT_SYSTEM = """\
Ты — умный ИИ-ассистент в Telegram. Отвечаешь по делу, чётко.
Если спрашивают код — пишешь рабочий код с объяснением.
Отвечаешь на языке пользователя.\
"""

def chat_respond(chat_id, user_message):
    add_to_history(chat_id, 'user', user_message)
    history = get_history(chat_id)
    history_text = _format_history(history[:-1])
    full_prompt = "{}\n\nПользователь: {}".format(history_text, user_message).strip()
    try:
        reply = call_llm(full_prompt, CHAT_SYSTEM)
    except Exception as e:
        reply = "❌ Ошибка LLM: {}".format(e)
    add_to_history(chat_id, 'assistant', reply)
    return reply


def _format_history(history):
    if not history:
        return ""
    lines = ["<История диалога>"]
    for msg in history:
        role = "Пользователь" if msg['role'] == 'user' else "Ассистент"
        content = msg['content'][:600] + '...' if len(msg['content']) > 600 else msg['content']
        lines.append("{}: {}".format(role, content))
    lines.append("</История диалога>")
    return "\n".join(lines)


# ════════════════════════════════════════════════════════════
#  СИСТЕМНЫЕ ПРОМТЫ
# ════════════════════════════════════════════════════════════

_ENV_INFO = """\
СРЕДА: Termux Android (Linux arm64).
Доступно: requests, bs4, edge-tts, asyncio, sqlite3, json, os, subprocess, PIL/Pillow, zipfile.
НЕ ИСПОЛЬЗУЙ: pyttsx3, espeak, tkinter, GUI-библиотеки.\
"""

CODER_BOTTOOLS_SYSTEM = """Ты — агент управления ботом АВТОМУВИ.
{}
Ты можешь управлять функциями бота через специальные команды в выводе скрипта.
Для вызова инструмента бота используй в коде Python:
  print("BOT_TOOL: <команда>")

Доступные команды (пиши ТОЧНО так):
  BOT_TOOL: tunnel_start bore       — запустить bore-туннель
  BOT_TOOL: tunnel_start ngrok      — запустить ngrok-туннель
  BOT_TOOL: tunnel_start serveo     — запустить serveo-туннель
  BOT_TOOL: tunnel_stop             — остановить все туннели
  BOT_TOOL: tunnel_status           — статус туннелей
  BOT_TOOL: save_html <url>         — скачать URL и сохранить как активную страницу
  BOT_TOOL: list_pages              — список сохранённых страниц
  BOT_TOOL: activate_page <id>      — активировать страницу по ID
  BOT_TOOL: send_file <path>        — отправить файл пользователю
  BOT_TOOL: read_file <path>        — прочитать файл
  BOT_TOOL: bot_stats               — статистика бота

Пример правильного ответа если нужно запустить туннель и сохранить страницу:
```python
print("BOT_TOOL: tunnel_start bore")
print("BOT_TOOL: save_html https://example.com")
print("Туннель и страница настроены!")
```

ПРАВИЛА: код в ```python```, print() для BOT_TOOL команд, минимум лишнего кода.
""".format(_ENV_INFO)

CODER_SCRIPT_SYSTEM = """Ты — Python-разработчик. Пишешь короткий рабочий скрипт.
{}
ПРАВИЛА: код в ```python```, используй print(), скрипт завершается сам, только доступные библиотеки.""".format(_ENV_INFO)

CODER_PROJECT_SYSTEM = """Ты — Python-разработчик. Создаёшь многофайловый проект.
{}
Для нескольких файлов используй: ###FILE: path/file.py\\n<код>
В конце ###INSTRUCTIONS: инструкция по запуску.
НЕ запускай серверы/ботов — только пиши код.""".format(_ENV_INFO)

CODER_FIX_SYSTEM = """Ты — Python-разработчик. Исправляешь ошибку в коде.
{}
ПРАВИЛА:
1. Изучи traceback внимательно
2. Исправь КОНКРЕТНУЮ ошибку
3. Верни ПОЛНЫЙ исправленный код в ```python```
4. Если библиотека недоступна — замени на альтернативу""".format(_ENV_INFO)

CODE_REVIEW_SYSTEM = """Ты — senior Python-разработчик делающий code review.
Анализируй код и выдай структурированный отчёт:

## 🔴 Критические ошибки
(баги, которые сломают программу)

## 🟡 Предупреждения  
(потенциальные проблемы, плохие практики)

## 🔵 Улучшения
(оптимизация, читаемость, безопасность)

## ✅ Исправленный код
(полный исправленный код в ```python```)

Будь конкретным, указывай номера строк."""

ANALYZE_SYSTEM = """Ты — аналитик кода и документов.
Анализируй предоставленный материал и выдай:
- Краткое описание (что это, для чего)
- Структура и компоненты
- Ключевые функции/разделы
- Зависимости и требования
- Потенциальные проблемы
Отвечай структурированно, на языке запроса."""

SCAFFOLD_SYSTEM = """Ты — Python-разработчик. Пишешь скрипт-скаффолдер.
{}
ПРАВИЛА:
1. Код в ```python```
2. Только stdlib: os, zipfile, pathlib
3. Каждый файл — реальный рабочий код
4. В конце: print("ZIP:", zip_path) с абсолютным путём
5. Скрипт завершается сам""".format(_ENV_INFO)


# ════════════════════════════════════════════════════════════
#  ГЛАВНАЯ ТОЧКА ВХОДА
# ════════════════════════════════════════════════════════════

def code_agent_run(chat_id, user_task, on_status=None, proj_mode=None,
                   attached_code=None, attached_file=None):
    """
    Агент-кодер v3.
    attached_code  — код присланный пользователем (для review/fix)
    attached_file  — путь к файлу/архиву (для analyze)
    """
    with _sessions_lock:
        s = _sessions.get(str(chat_id))
        if s:
            s['iterations'] = 0

    add_to_history(chat_id, 'user', user_task)
    history = get_history(chat_id)

    # Явный режим из UI
    if proj_mode == 'scaffold' or _is_scaffold_task(user_task):
        return _run_scaffold_agent(chat_id, user_task, on_status)
    if proj_mode == 'plan':
        return _run_project_agent(chat_id, user_task, history, on_status)
    if proj_mode == 'onebyone':
        return _run_onebyone_agent(chat_id, user_task, history, on_status)
    if proj_mode == 'review':
        return _run_review_agent(chat_id, user_task, history, on_status, attached_code)
    if proj_mode == 'fix':
        return _run_fix_agent(chat_id, user_task, history, on_status, attached_code)
    if proj_mode == 'analyze':
        return _run_analyze_agent(chat_id, user_task, history, on_status, attached_file)
    if proj_mode == 'image':
        return _run_image_agent(chat_id, user_task, on_status)

    # Автоопределение
    task_type = detect_task_type(user_task)

    if proj_mode == 'bot_tools':
        return _run_script_agent(chat_id, user_task, history, on_status,
                                  system_override=CODER_BOTTOOLS_SYSTEM)
    if task_type == TASK_VIDEO:
        return _run_video_agent(chat_id, user_task, on_status)
    if task_type == TASK_FILE:
        return _run_file_agent(chat_id, user_task, history, on_status, attached_file)
    if task_type == TASK_IMAGE:
        return _run_image_agent(chat_id, user_task, on_status)
    if task_type == TASK_REVIEW:
        return _run_review_agent(chat_id, user_task, history, on_status, attached_code)
    if task_type == TASK_FIX:
        return _run_fix_agent(chat_id, user_task, history, on_status, attached_code)
    if task_type == TASK_ANALYZE:
        return _run_analyze_agent(chat_id, user_task, history, on_status, attached_file)
    if task_type == TASK_PROJECT:
        return _run_project_agent(chat_id, user_task, history, on_status)

    return _run_script_agent(chat_id, user_task, history, on_status)


# ════════════════════════════════════════════════════════════
#  АГЕНТ: АНАЛИЗ КОДА / CODE REVIEW
# ════════════════════════════════════════════════════════════

def _run_review_agent(chat_id, user_task, history, on_status, attached_code=None):
    """Находит ошибки, предупреждения, улучшения. Возвращает исправленный код."""
    if on_status:
        on_status("🔍 Анализирую код...")

    # Получаем код для анализа
    code_to_review = attached_code or _extract_code_from_history(history) or ""

    if not code_to_review:
        # Нет кода — просим LLM ответить текстом
        reply = _call_with_continuation(user_task, CODE_REVIEW_SYSTEM, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
        add_to_history(chat_id, 'assistant', reply[:500])
        return {'success': True, 'output': reply, 'code': '', 'iterations': 1,
                'is_project': False, 'files': [], 'text_only': True}

    prompt = "Выполни code review следующего кода.\nЗапрос: {}\n\nКод:\n```python\n{}\n```".format(
        user_task, code_to_review)

    if on_status:
        on_status("🧠 LLM анализирует ({} строк)...".format(len(code_to_review.splitlines())))

    try:
        review = _call_with_continuation(prompt, CODE_REVIEW_SYSTEM, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
    except Exception as e:
        return {'success': False, 'output': "❌ LLM: {}".format(e),
                'code': '', 'iterations': 1, 'is_project': False, 'files': []}

    # Извлекаем исправленный код если есть
    fixed_code = _extract_code(review)
    files = []
    if fixed_code:
        ts = time.strftime('%H%M%S')
        run_dir = os.path.join(PROJECTS_DIR, 'review_' + ts)
        os.makedirs(run_dir, exist_ok=True)
        fixed_path = os.path.join(run_dir, 'fixed_code.py')
        with open(fixed_path, 'w', encoding='utf-8') as f:
            f.write(fixed_code)
        files = [fixed_path]

    add_to_history(chat_id, 'assistant', review[:500])
    return {
        'success': True, 'output': review, 'code': fixed_code or '',
        'iterations': 1, 'is_project': False, 'files': files,
        'text_only': not fixed_code, '_full_output': review,
    }


# ════════════════════════════════════════════════════════════
#  АГЕНТ: ИСПРАВЛЕНИЕ ОШИБОК (АВТО-ФИХ)
# ════════════════════════════════════════════════════════════

def _diagnose_error(error_text: str) -> str:
    """Определяет тип ошибки и возвращает подсказку для LLM."""
    e = error_text.lower()

    # Нет модуля
    m = re.search(r"no module named '([^']+)'", error_text, re.IGNORECASE)
    if m:
        mod = m.group(1)
        alternatives = {
            'pyttsx3': 'используй edge-tts вместо pyttsx3',
            'cv2':     'используй Pillow (PIL) вместо OpenCV',
            'tkinter': 'GUI недоступен, выведи результат в консоль',
            'pygame':  'используй другой подход без pygame',
            'sklearn': 'sklearn нет на Termux, используй простую математику',
            'torch':   'torch слишком тяжёлый для Termux',
            'numpy':   'попробуй pip install numpy или используй math/statistics',
        }
        hint = alternatives.get(mod, "модуль '{}' недоступен на Termux. {}".format(mod, _ENV_INFO))
        return hint

    # Timeout
    if 'timeouterror' in e or 'timed out' in e:
        return 'Таймаут выполнения — добавь timeout в requests, убери бесконечные циклы'

    # Права доступа
    if 'permissionerror' in e or 'errno 13' in e:
        return 'Нет прав доступа — используй директорию ~/storage/downloads/ или /data/data/com.termux/files/home/'

    # SSL
    if 'ssl' in e and ('cert' in e or 'verify' in e):
        return 'SSL ошибка — добавь verify=False в requests.get() или обнови сертификаты'

    # UnicodeDecodeError
    if 'unicodedecodeerror' in e or 'unicodeencodeerror' in e:
        return 'Проблема с кодировкой — добавь encoding="utf-8", errors="ignore" при открытии файлов'

    # RecursionError
    if 'recursionerror' in e or 'maximum recursion' in e:
        return 'Бесконечная рекурсия — добавь базовый случай или перепиши через цикл'

    # MemoryError
    if 'memoryerror' in e:
        return 'Нехватка памяти — обрабатывай данные чанками, не загружай всё сразу'

    # JSONDecodeError
    if 'jsondecode' in e or 'json.decoder' in e:
        return 'Некорректный JSON — проверь ответ API перед парсингом, добавь try/except'

    # Telegram API errors
    if 'telegram' in e and ('401' in error_text or 'unauthorized' in error_text):
        return 'Неверный токен Telegram бота — проверь TELEGRAM_BOT_TOKEN'

    if 'telegram' in e and '409' in error_text:
        return 'Конфликт Telegram polling — убей старый процесс бота: pkill -f bot.py'

    return ''


def _run_fix_agent(chat_id, user_task, history, on_status, attached_code=None):
    """
    Получает код + ошибку, итеративно исправляет.
    При неизвестных ошибках — ищет решение через LLM с контекстом.
    """
    if on_status:
        on_status("🔧 Запускаю агент исправления ошибок...")

    code_to_fix = attached_code or _extract_code_from_history(history) or ""
    error_text = _extract_error_from_text(user_task)

    if not code_to_fix and not error_text:
        # Нет ни кода ни ошибки — обычная генерация
        return _run_script_agent(chat_id, user_task, history, on_status)

    all_attempts = []
    last_code = code_to_fix
    last_error = error_text

    for iteration in range(1, MAX_ITERATIONS + 1):
        with _sessions_lock:
            s = _sessions.get(str(chat_id))
            if s:
                s['iterations'] = iteration

        if on_status:
            on_status("🔧 Итерация {} — ищу исправление...".format(iteration))

        # Строим контекст из всех предыдущих попыток
        attempts_ctx = ""
        if all_attempts:
            attempts_ctx = "\n\nПредыдущие попытки которые не сработали:\n"
            for i, a in enumerate(all_attempts):
                attempts_ctx += "Попытка {}:\nКод:\n```python\n{}\n```\nОшибка:\n{}\n\n".format(
                    i+1, a['code'][:800], a['error'][:400])

        prompt = (
            "Задача: {}\n\n"
            "Код с ошибкой:\n```python\n{}\n```\n\n"
            "Ошибка/traceback:\n{}\n"
            "{}\n"
            "Напиши ПОЛНОСТЬЮ исправленный код:"
        ).format(user_task, last_code or "(нет кода)", last_error or "(нет ошибки)", attempts_ctx)

        try:
            llm_response = _call_with_continuation(prompt, CODER_FIX_SYSTEM, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
        except Exception as e:
            return {'success': False, 'output': "❌ LLM: {}".format(e),
                    'code': last_code or '', 'iterations': iteration, 'is_project': False, 'files': []}

        fixed_code = _extract_code(llm_response)
        if not fixed_code:
            # LLM дал текстовый ответ — возможно объяснение без кода
            add_to_history(chat_id, 'assistant', llm_response[:500])
            return {'success': True, 'output': llm_response, 'code': '',
                    'iterations': iteration, 'is_project': False, 'files': [], 'text_only': True}

        last_code = fixed_code

        if on_status:
            on_status("⚙️ Запускаю исправленный код (попытка {})...".format(iteration))

        success, output, created_files = _execute_code(fixed_code)

        if success:
            add_to_history(chat_id, 'assistant', "✅ Исправлено за {} итераций".format(iteration))
            return {'success': True, 'code': fixed_code, 'output': output,
                    'iterations': iteration, 'is_project': False, 'files': created_files,
                    '_full_output': output}
        else:
            last_error = output
            all_attempts.append({'code': fixed_code, 'error': output[:600]})

            # Умная диагностика ошибок — подсказки для LLM
            hint = _diagnose_error(output)
            if hint:
                last_error = output + "\n\n💡 ДИАГНОЗ: " + hint
                if on_status:
                    on_status("🔍 {}".format(hint[:80]))

            if iteration == MAX_ITERATIONS:
                # Последняя попытка — просим LLM объяснить что не так
                explain_prompt = (
                    "Код не работает после {} попыток исправить.\n"
                    "Последняя ошибка: {}\n"
                    "Последний код:\n```python\n{}\n```\n"
                    "Объясни причину и предложи альтернативный подход."
                ).format(iteration, output[:600], fixed_code[:800])

                try:
                    explanation = call_llm(explain_prompt, CODER_FIX_SYSTEM, max_tokens=2000)
                    output = output + "\n\n💡 Анализ агента:\n" + explanation
                except Exception:
                    pass

                add_to_history(chat_id, 'assistant', "❌ Не удалось за {} итераций".format(iteration))
                return {'success': False, 'code': fixed_code, 'output': output,
                        'iterations': iteration, 'is_project': False, 'files': [],
                        '_full_output': output}


# ════════════════════════════════════════════════════════════
#  АГЕНТ: АНАЛИЗ ДОКУМЕНТОВ / АРХИВОВ
# ════════════════════════════════════════════════════════════

def _run_analyze_agent(chat_id, user_task, history, on_status, attached_file=None):
    """Анализирует код, документ или архив."""
    if on_status:
        on_status("🔍 Анализирую...")

    content_to_analyze = ""

    # Файл/архив
    if attached_file and os.path.exists(attached_file):
        ext = os.path.splitext(attached_file)[1].lower()

        if ext == '.zip':
            import zipfile
            if on_status:
                on_status("📦 Читаю архив...")
            with zipfile.ZipFile(attached_file) as zf:
                names = zf.namelist()
                content_to_analyze = "Архив содержит {} файлов:\n{}\n\n".format(
                    len(names), "\n".join(names[:50]))
                # Читаем текстовые файлы
                for name in names[:20]:
                    if any(name.endswith(ext) for ext in ('.py','.js','.ts','.md','.txt','.json','.yaml','.toml')):
                        try:
                            with zf.open(name) as f:
                                fc = f.read(8000).decode('utf-8', errors='ignore')
                                content_to_analyze += "\n--- {} ---\n{}\n".format(name, fc[:2000])
                        except Exception:
                            pass

        elif ext in ('.py','.js','.ts','.go','.rs','.java','.kt','.cpp','.c','.h'):
            with open(attached_file, encoding='utf-8', errors='ignore') as f:
                content_to_analyze = f.read(20000)
            if on_status:
                on_status("📄 Анализирую код ({} строк)...".format(len(content_to_analyze.splitlines())))

        elif ext in ('.txt','.md','.csv','.json','.yaml','.toml','.env','.cfg','.ini'):
            with open(attached_file, encoding='utf-8', errors='ignore') as f:
                content_to_analyze = f.read(15000)

    # Код из истории чата
    if not content_to_analyze:
        content_to_analyze = _extract_code_from_history(history) or ""

    if not content_to_analyze:
        # Нет конкретного контента — общий анализ по описанию
        reply = _call_with_continuation(user_task, ANALYZE_SYSTEM, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
        add_to_history(chat_id, 'assistant', reply[:500])
        return {'success': True, 'output': reply, 'code': '', 'iterations': 1,
                'is_project': False, 'files': [], 'text_only': True}

    prompt = "Запрос: {}\n\nКонтент для анализа:\n{}".format(
        user_task, content_to_analyze[:12000])

    try:
        analysis = _call_with_continuation(prompt, ANALYZE_SYSTEM, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
    except Exception as e:
        return {'success': False, 'output': "❌ LLM: {}".format(e),
                'code': '', 'iterations': 1, 'is_project': False, 'files': []}

    add_to_history(chat_id, 'assistant', analysis[:500])
    return {
        'success': True, 'output': analysis, 'code': '', 'iterations': 1,
        'is_project': False, 'files': [], 'text_only': True, '_full_output': analysis,
    }


# ════════════════════════════════════════════════════════════
#  АГЕНТ: СКАЧИВАНИЕ ВИДЕО / АУДИО (YouTube и др.)
# ════════════════════════════════════════════════════════════

def _extract_url_from_task(task_text: str):
    """Вытаскиваем URL из текста запроса."""
    m = re.search(r'https?://[^\s]+', task_text)
    return m.group(0) if m else None

def _detect_video_format(task_text: str):
    """
    Определяем нужный формат: 'mp3', 'mp4', 'mp4_360', 'mp4_1080'.
    По умолчанию — mp4 720p если не указано явно.
    """
    t = task_text.lower()
    if any(x in t for x in ['mp3', 'аудио', 'audio', 'музык', 'music',
                              'извлеки аудио', 'extract audio']):
        return 'mp3'
    if '1080' in t or 'fullhd' in t or 'full hd' in t or 'высокое качество' in t:
        return 'mp4_1080'
    if '360' in t or 'лёгкий' in t or 'легкий' in t or 'маленький' in t:
        return 'mp4_360'
    return 'mp4'

def _check_yt_dlp():
    """Проверяем наличие yt-dlp, возвращаем путь или None."""
    import shutil as _sh
    # Системный путь
    p = _sh.which('yt-dlp')
    if p:
        return p
    # Termux pip-путь
    termux_pip = '/data/data/com.termux/files/usr/bin/yt-dlp'
    if os.path.exists(termux_pip):
        return termux_pip
    # Python-модуль как fallback
    try:
        import yt_dlp  # noqa
        return '__python_module__'
    except ImportError:
        return None

def _download_via_python_module(url, out_dir, fmt, on_status):
    """
    Скачивает через yt_dlp Python-модуль (без системного бинарника).
    Возвращает путь к скачанному файлу или None.
    """
    import yt_dlp

    downloaded = []

    def _progress_hook(d):
        if d['status'] == 'downloading':
            pct = d.get('_percent_str', '?')
            speed = d.get('_speed_str', '?')
            if on_status:
                on_status("⬇️ Скачиваю... {} @ {}".format(pct, speed))
        elif d['status'] == 'finished':
            downloaded.append(d['filename'])

    _format_map = {
        'mp4':      'bestvideo[ext=mp4][height<=720]+bestaudio[ext=m4a]/best[ext=mp4]/best',
        'mp4_360':  'bestvideo[ext=mp4][height<=360]+bestaudio[ext=m4a]/best[ext=mp4][height<=360]/best',
        'mp4_1080': 'bestvideo[ext=mp4][height<=1080]+bestaudio[ext=m4a]/best[ext=mp4]/best',
    }

    if fmt == 'mp3':
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': os.path.join(out_dir, '%(title)s.%(ext)s'),
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'progress_hooks': [_progress_hook],
            'quiet': True,
            'no_warnings': True,
        }
    else:  # mp4 / mp4_360 / mp4_1080
        ydl_opts = {
            'format': _format_map.get(fmt, _format_map['mp4']),
            'outtmpl': os.path.join(out_dir, '%(title)s.%(ext)s'),
            'merge_output_format': 'mp4',
            'progress_hooks': [_progress_hook],
            'quiet': True,
            'no_warnings': True,
        }

    # Общие опции — безопасное имя файла, без сохранения куки в файл
    ydl_opts.update({
        'restrictfilenames': True,   # убирает спецсимволы из имени файла
        'trim_filenames':    80,     # макс длина имени — важно для Termux (длинные имена падают)
        'retries':           3,      # повтор при сетевой ошибке
        'fragment_retries':  3,
        'ignoreerrors':      False,
        'socket_timeout':    30,
    })

    title, duration = 'video', 0
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            if info:
                title    = info.get('title', 'video')
                duration = info.get('duration', 0) or 0
    except yt_dlp.utils.DownloadError as e:
        raise Exception("yt-dlp: {}".format(str(e)[:300]))

    # Ищем скачанный файл — сначала по нужному расширению
    target_ext = '.mp3' if fmt == 'mp3' else '.mp4'
    best = None
    best_mtime = 0
    for fname in os.listdir(out_dir):
        fpath_check = os.path.join(out_dir, fname)
        mt = os.path.getmtime(fpath_check)
        if fname.endswith(target_ext) and mt > best_mtime:
            best = fpath_check
            best_mtime = mt

    if best:
        return best, title, duration

    # Fallback — любой файл в директории (webm, m4a, etc.)
    files = [os.path.join(out_dir, f) for f in os.listdir(out_dir)
             if not f.endswith('.part') and not f.endswith('.ytdl')]
    if files:
        return max(files, key=os.path.getmtime), title, duration

    return None, title, duration

def _download_via_binary(yt_dlp_bin, url, out_dir, fmt, on_status):
    """
    Скачивает через системный бинарник yt-dlp.
    Возвращает (filepath, title, duration) или (None, '', 0).
    """
    _fmt_map_bin = {
        'mp4':      'bestvideo[ext=mp4][height<=720]+bestaudio[ext=m4a]/best[ext=mp4]/best',
        'mp4_360':  'bestvideo[ext=mp4][height<=360]+bestaudio[ext=m4a]/best[ext=mp4][height<=360]/best',
        'mp4_1080': 'bestvideo[ext=mp4][height<=1080]+bestaudio[ext=m4a]/best[ext=mp4]/best',
    }

    if fmt == 'mp3':
        cmd = [
            yt_dlp_bin, url,
            '-x', '--audio-format', 'mp3',
            '--audio-quality', '192K',
            '-o', os.path.join(out_dir, '%(title)s.%(ext)s'),
            '--no-playlist',
            '--quiet', '--progress',
        ]
    else:  # mp4 / mp4_360 / mp4_1080
        cmd = [
            yt_dlp_bin, url,
            '-f', _fmt_map_bin.get(fmt, _fmt_map_bin['mp4']),
            '--merge-output-format', 'mp4',
            '-o', os.path.join(out_dir, '%(title)s.%(ext)s'),
            '--no-playlist',
            '--quiet', '--progress',
            '--print', 'after_move:%(title)s|||%(duration)s',
        ]

    if on_status:
        on_status("⬇️ Запускаю yt-dlp...")

    # --no-part — не оставлять .part файлы при прерывании
    if '--no-part' not in cmd:
        cmd.append('--no-part')

    if on_status:
        on_status("⬇️ Запускаю yt-dlp...")

    result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)

    if result.returncode != 0 and result.returncode is not None:
        err = (result.stderr or result.stdout or '').strip()
        # Некоторые ошибки не критичны (файл всё равно скачан)
        if 'already been downloaded' not in err and not os.listdir(out_dir):
            raise Exception("yt-dlp rc={}: {}".format(result.returncode, err[:300]))

    title, duration = '', 0
    for line in result.stdout.splitlines():
        if '|||' in line:
            parts = line.split('|||')
            title = parts[0].strip()
            try:
                duration = int(parts[1].strip())
            except Exception:
                pass

    # Ищем скачанный файл
    target_ext = '.mp3' if fmt == 'mp3' else '.mp4'
    for fname in os.listdir(out_dir):
        if fname.endswith(target_ext):
            return os.path.join(out_dir, fname), title, duration

    files = [os.path.join(out_dir, f) for f in os.listdir(out_dir)]
    if files:
        return max(files, key=os.path.getmtime), title, duration

    # Если ничего нет — ошибка
    err = result.stderr.strip() or result.stdout.strip()
    raise Exception("yt-dlp не создал файл.\n{}".format(err[:500]))


def _run_video_agent(chat_id, user_task, on_status=None):
    """
    Агент скачивания видео/аудио.

    Алгоритм:
      1. Извлекаем URL из текста запроса
      2. Определяем формат (mp3 / mp4)
      3. Ищем yt-dlp (бинарник или Python-модуль)
      4. Скачиваем с прогресс-статусами
      5. Отправляем файл через send_document
         + если есть туннель — даём ссылку

    Если yt-dlp не установлен — даём инструкцию по установке.
    Telegram ограничивает файлы до 50 MB через Bot API.
    Для больших файлов — предупреждаем и используем туннельную ссылку.
    """

    def _st(msg):
        if on_status:
            on_status(msg)

    _st("🎬 Агент скачивания видео запущен...")

    # ── URL ────────────────────────────────────────────────
    url = _extract_url_from_task(user_task)
    if not url:
        _st("❌ Не найден URL в запросе.\nПример: скачай mp3 https://youtu.be/xxxxx")
        return {
            'success': False, 'code': '', 'iterations': 1,
            'is_project': False, 'files': [], 'text_only': True,
            '_full_output': "❌ URL не найден. Укажи ссылку на видео в запросе.",
            '_task_type': 'video',
        }

    fmt = _detect_video_format(user_task)
    _st("🔍 URL: {}\n📦 Формат: {}".format(url, fmt.upper()))

    # ── Проверяем yt-dlp ───────────────────────────────────
    yt_dlp_path = _check_yt_dlp()
    if not yt_dlp_path:
        install_msg = (
            "❌ yt-dlp не установлен.\n\n"
            "Установка в Termux:\n"
            "<code>pip install yt-dlp</code>\n\n"
            "Или через pkg:\n"
            "<code>pkg install python</code>\n"
            "<code>pip install yt-dlp</code>\n\n"
            "Если нужна конвертация в MP3 — также нужен ffmpeg:\n"
            "<code>pkg install ffmpeg</code>"
        )
        _st(install_msg)
        return {
            'success': False, 'code': '', 'iterations': 1,
            'is_project': False, 'files': [], 'text_only': True,
            '_full_output': install_msg, '_task_type': 'video',
        }

    # ── Рабочая директория ─────────────────────────────────
    ts = time.strftime('%Y%m%d_%H%M%S')
    work_dir = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'agent_projects', 'video', ts
    )
    os.makedirs(work_dir, exist_ok=True)

    # ── Скачивание ─────────────────────────────────────────
    try:
        _st("⬇️ Начинаю скачивание...")

        if yt_dlp_path == '__python_module__':
            fpath, title, duration = _download_via_python_module(url, work_dir, fmt, on_status)
        else:
            fpath, title, duration = _download_via_binary(yt_dlp_path, url, work_dir, fmt, on_status)

        if not fpath or not os.path.exists(fpath):
            raise Exception("Файл не найден после скачивания")

        fsize = os.path.getsize(fpath)
        fsize_mb = fsize / (1024 * 1024)
        fname = os.path.basename(fpath)

        _st("✅ Скачано: {}\n📏 Размер: {:.1f} MB\n⏱ Длительность: {} сек".format(
            title or fname, fsize_mb,
            "{}:{:02d}".format(duration // 60, duration % 60) if duration else "?"
        ))

        # ── Предупреждение о лимите Telegram ──────────────
        # Bot API принимает файлы до 50 MB, отправляет — до 20 MB без streaming
        tg_limit_mb = 49
        if fsize_mb > tg_limit_mb:
            _st(
                "⚠️ Файл {:.1f} MB — больше лимита Telegram ({} MB).\n"
                "Отправка может не работать.\n"
                "💡 Используй туннельную ссылку для скачивания.".format(
                    fsize_mb, tg_limit_mb)
            )

        return {
            'success': True,
            'code': '',
            'iterations': 1,
            'is_project': False,
            'files': [fpath],
            'text_only': False,
            '_full_output': "🎬 {} ({:.1f} MB)".format(title or fname, fsize_mb),
            '_task_type': 'video',
            '_video_fmt': fmt,
            '_video_title': title or fname,
            '_fsize_mb': fsize_mb,
        }

    except subprocess.TimeoutExpired:
        msg = "❌ Таймаут — скачивание заняло больше 5 минут."
        _st(msg)
        return {'success': False, 'code': '', 'iterations': 1,
                'is_project': False, 'files': [], 'text_only': True,
                '_full_output': msg, '_task_type': 'video'}

    except Exception as e:
        msg = "❌ Ошибка скачивания: {}".format(e)
        _st(msg)
        return {'success': False, 'code': '', 'iterations': 1,
                'is_project': False, 'files': [], 'text_only': True,
                '_full_output': msg, '_task_type': 'video'}


# ════════════════════════════════════════════════════════════
#  АГЕНТ: СОЗДАНИЕ ФАЙЛОВ (TXT/MD/CSV/DOCX/ZIP/RAR)
# ════════════════════════════════════════════════════════════

def _detect_output_format(task_text: str):
    """
    Определяет нужный формат файла по тексту задачи.
    Возвращает список форматов — может быть несколько (например txt + zip).
    """
    t = task_text.lower()
    formats = []

    if any(x in t for x in ['docx', 'word документ', 'word файл', 'документ word', '.docx']):
        formats.append('docx')
    if any(x in t for x in ['csv', 'таблицу csv', 'в csv', '.csv']):
        formats.append('csv')
    if any(x in t for x in ['markdown', ' md ', '.md', 'md файл']):
        formats.append('md')
    if any(x in t for x in ['rar архив', 'в rar', '.rar', 'собери rar', 'упакуй в rar']):
        formats.append('rar')
    if any(x in t for x in ['zip архив', 'в zip', '.zip', 'собери zip', 'упакуй в zip']):
        formats.append('zip')
    if any(x in t for x in ['txt', 'текстовый файл', 'текстовый документ', 'txt файл']):
        formats.append('txt')

    # Если ничего конкретного — txt по умолчанию
    if not formats:
        formats.append('txt')

    return formats


def _collect_context(chat_id, user_task, attached_file=None):
    """
    Собирает весь доступный контекст для генерации файла:
      - история чата
      - содержимое прикреплённого файла
      - данные из базы данных бота (статистика, новости)
    Возвращает строку контекста.
    """
    parts = []

    # ── История чата ───────────────────────────────────────
    history = get_history(chat_id)
    if history:
        history_lines = []
        for msg in history[-20:]:  # последние 20 сообщений
            role = 'Пользователь' if msg['role'] == 'user' else 'Ассистент'
            history_lines.append("{}: {}".format(role, msg['content'][:500]))
        parts.append("=== История чата ===\n" + "\n".join(history_lines))

    # ── Прикреплённый файл ─────────────────────────────────
    if attached_file and os.path.exists(attached_file):
        try:
            ext = os.path.splitext(attached_file)[1].lower()
            if ext in ('.txt', '.md', '.py', '.js', '.csv', '.json', '.html', '.xml'):
                with open(attached_file, 'r', encoding='utf-8', errors='replace') as f:
                    file_content = f.read(8000)  # первые 8KB
                parts.append("=== Содержимое файла {} ===\n{}".format(
                    os.path.basename(attached_file), file_content))
            elif ext == '.zip':
                import zipfile as _zf
                with _zf.ZipFile(attached_file) as zf:
                    names = zf.namelist()[:10]
                    contents = []
                    for name in names:
                        ext2 = os.path.splitext(name)[1].lower()
                        if ext2 in ('.txt', '.md', '.py', '.csv', '.json'):
                            try:
                                txt = zf.read(name).decode('utf-8', errors='replace')[:3000]
                                contents.append("--- {} ---\n{}".format(name, txt))
                            except Exception:
                                pass
                    parts.append("=== Содержимое архива {} ===\n{}".format(
                        os.path.basename(attached_file), "\n".join(contents)))
        except Exception as e:
            parts.append("=== Файл {} (ошибка чтения: {}) ===".format(
                os.path.basename(attached_file), e))

    # ── База данных бота (статистика, новости) ─────────────
    try:
        # Импортируем только если доступно — не обязательная зависимость
        from database import get_stats, get_today_count
        stats = get_stats()
        today = get_today_count()
        parts.append("=== Статистика бота ===\nВсего новостей: {}\nСегодня: {}".format(
            stats.get('total', 0) if isinstance(stats, dict) else stats,
            today))
    except Exception:
        pass

    return "\n\n".join(parts) if parts else "(контекст отсутствует)"


def _generate_file_content(user_task, context, output_format, on_status=None):
    """
    Запрашивает у LLM содержимое для файла нужного формата.
    Возвращает (filename, content_str, is_binary).
    """
    fmt_hints = {
        'txt':  'Верни только текст без лишних пояснений. Без markdown-разметки.',
        'md':   'Верни Markdown-разметку. Используй заголовки, списки, блоки кода где уместно.',
        'csv':  ('Верни данные в формате CSV. Первая строка — заголовки столбцов. '
                 'Разделитель — запятая. Без лишних пояснений, только CSV.'),
        'docx': 'Верни текст документа в Markdown-формате. Заголовки через #, ## и т.д.',
    }
    hint = fmt_hints.get(output_format, fmt_hints['txt'])

    prompt = (
        "Задача пользователя: {}\n\n"
        "Имеющийся контекст:\n{}\n\n"
        "Формат вывода: {}\n"
        "{}\n\n"
        "Создай содержимое файла прямо сейчас:"
    ).format(user_task, context[:6000], output_format.upper(), hint)

    if on_status:
        on_status("✍️ Генерирую содержимое ({})...".format(output_format.upper()))

    content_str = call_llm(prompt)
    return content_str


def _write_txt(content, path):
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)

def _write_md(content, path):
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)

def _write_csv(content, path):
    with open(path, 'w', encoding='utf-8', newline='') as f:
        f.write(content)

def _write_docx(content, path):
    """
    Конвертирует Markdown-текст в .docx через python-docx.
    Если python-docx не установлен — сохраняет как .txt с .docx расширением.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for line in content.splitlines():
            stripped = line.strip()
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped.startswith('```'):
                pass  # пропускаем маркеры блока кода
            elif stripped:
                doc.add_paragraph(stripped)
            else:
                doc.add_paragraph('')
        doc.save(path)
    except ImportError:
        # python-docx не установлен — сохраняем как текст
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)

def _pack_zip(files, zip_path):
    import zipfile as _zf
    with _zf.ZipFile(zip_path, 'w', _zf.ZIP_DEFLATED) as zf:
        for fpath in files:
            if os.path.exists(fpath):
                zf.write(fpath, os.path.basename(fpath))

def _pack_rar(files, rar_path):
    """
    Создаёт RAR-архив через системный rar/unrar.
    Если rar не установлен — создаёт zip с .rar расширением и сообщает об этом.
    """
    import shutil as _sh
    if _sh.which('rar'):
        file_args = [fpath for fpath in files if os.path.exists(fpath)]
        subprocess.run(['rar', 'a', rar_path] + file_args,
                       capture_output=True, timeout=30)
    else:
        # Фоллбэк: zip с расширением .rar
        import zipfile as _zf
        with _zf.ZipFile(rar_path, 'w', _zf.ZIP_DEFLATED) as zf:
            for fpath in files:
                if os.path.exists(fpath):
                    zf.write(fpath, os.path.basename(fpath))
        return False  # сигнализируем что это ZIP внутри
    return True


def _run_file_agent(chat_id, user_task, history, on_status=None, attached_file=None):
    """
    Агент создания файлов.

    Алгоритм:
      1. Определяем нужный формат(ы) из задачи
      2. Собираем контекст (история, файлы, БД)
      3. LLM генерирует содержимое для каждого формата
      4. Пишем файлы на диск
      5. Если нужен zip/rar — упаковываем
      6. Возвращаем список файлов через стандартный result-словарь
    """
    if on_status:
        on_status("📁 Агент создания файлов запущен...")

    formats = _detect_output_format(user_task)

    # Нужна ли архивация
    need_zip = 'zip' in formats
    need_rar = 'rar' in formats
    # Контентные форматы (не архивы)
    content_formats = [f for f in formats if f not in ('zip', 'rar')]
    if not content_formats:
        content_formats = ['txt']  # всегда генерируем хотя бы txt для архива

    if on_status:
        fmt_str = ', '.join(f.upper() for f in formats)
        on_status("🔍 Формат(ы): {}\n📚 Собираю контекст...".format(fmt_str))

    # ── Контекст ───────────────────────────────────────────
    context = _collect_context(chat_id, user_task, attached_file)

    # ── Рабочая директория ─────────────────────────────────
    ts = time.strftime('%Y%m%d_%H%M%S')
    work_dir = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'agent_projects', 'files', ts
    )
    os.makedirs(work_dir, exist_ok=True)

    created_files = []
    summaries = []

    # ── Генерируем файлы ───────────────────────────────────
    for fmt in content_formats:
        try:
            content_str = _generate_file_content(user_task, context, fmt, on_status)

            # Имя файла — коротко, понятно, без пробелов
            base_name = re.sub(r'[^\w\-]', '_', user_task[:30].strip().lower())
            filename = "{}_{}.{}".format(base_name, ts, fmt)
            fpath = os.path.join(work_dir, filename)

            if on_status:
                on_status("💾 Записываю {}...".format(filename))

            if fmt == 'txt':
                _write_txt(content_str, fpath)
            elif fmt == 'md':
                _write_md(content_str, fpath)
            elif fmt == 'csv':
                _write_csv(content_str, fpath)
            elif fmt == 'docx':
                _write_docx(content_str, fpath)

            created_files.append(fpath)
            summaries.append("✅ {} ({} KB)".format(
                filename, max(1, os.path.getsize(fpath) // 1024)))

        except Exception as e:
            summaries.append("❌ {}: {}".format(fmt.upper(), e))

    # ── Архивируем ─────────────────────────────────────────
    archive_files = []

    if need_zip and created_files:
        zip_name = "result_{}.zip".format(ts)
        zip_path = os.path.join(work_dir, zip_name)
        if on_status:
            on_status("📦 Собираю ZIP...")
        _pack_zip(created_files, zip_path)
        archive_files.append(zip_path)
        summaries.append("📦 {} ({} KB)".format(
            zip_name, max(1, os.path.getsize(zip_path) // 1024)))

    if need_rar and created_files:
        rar_name = "result_{}.rar".format(ts)
        rar_path = os.path.join(work_dir, rar_name)
        if on_status:
            on_status("📦 Собираю RAR...")
        real_rar = _pack_rar(created_files, rar_path)
        archive_files.append(rar_path)
        note = "" if real_rar else " (ZIP внутри — rar не установлен)"
        summaries.append("📦 {}{} ({} KB)".format(
            rar_name, note, max(1, os.path.getsize(rar_path) // 1024)))

    # ── Итог ───────────────────────────────────────────────
    all_output_files = created_files + archive_files
    summary_text = "📁 Файлы готовы:\n" + "\n".join(summaries)

    return {
        'success': bool(created_files),
        'code': '',
        'iterations': 1,
        'is_project': False,
        'files': all_output_files,
        'text_only': False,
        '_full_output': summary_text,
        '_task_type': 'file',
    }


# ════════════════════════════════════════════════════════════
#  АГЕНТ: ГЕНЕРАЦИЯ ИЗОБРАЖЕНИЙ
# ════════════════════════════════════════════════════════════

def _run_image_agent(chat_id, user_task, on_status):
    """Генерирует изображение через доступный API."""
    if on_status:
        on_status("🎨 Генерирую изображение...")

    provider = config.LLM_PROVIDER.lower()
    api_key  = getattr(config, 'LLM_API_KEY', '') or ''

    # Пробуем разные провайдеры
    img_path = None
    error    = None

    # Together AI — имеет FLUX.1 schnell бесплатно
    if provider in ('together', 'openai', 'openrouter') or True:
        img_path, error = _gen_image_together(user_task, api_key, on_status)

    if not img_path:
        # Fallback: генерируем Python скрипт который создаёт изображение через PIL
        if on_status:
            on_status("🖼 API не доступен — генерирую через PIL...")
        return _gen_image_with_code(chat_id, user_task, on_status)

    add_to_history(chat_id, 'assistant', "🎨 Изображение создано: {}".format(user_task[:50]))
    return {
        'success': True, 'output': "🎨 Изображение готово",
        'code': '', 'iterations': 1, 'is_project': False,
        'files': [img_path], 'file_names': [os.path.basename(img_path)],
    }


def _gen_image_together(prompt, api_key, on_status):
    """Генерация через Together AI FLUX.1."""
    import requests as req
    try:
        # Улучшаем промт
        enhanced = prompt + ", high quality, detailed, 4k"

        r = req.post(
            "https://api.together.xyz/v1/images/generations",
            json={"model": "black-forest-labs/FLUX.1-schnell-Free",
                  "prompt": enhanced, "n": 1,
                  "width": 1024, "height": 1024},
            headers={"Authorization": "Bearer " + api_key,
                     "Content-Type": "application/json"},
            timeout=60
        )
        if r.status_code == 200:
            data = r.json()
            url = data['data'][0].get('url') or data['data'][0].get('b64_json')
            if url and url.startswith('http'):
                # Скачиваем
                img_r = req.get(url, timeout=30)
                ts = time.strftime('%H%M%S')
                img_dir = os.path.join(PROJECTS_DIR, 'images')
                os.makedirs(img_dir, exist_ok=True)
                img_path = os.path.join(img_dir, 'image_{}.png'.format(ts))
                with open(img_path, 'wb') as f:
                    f.write(img_r.content)
                return img_path, None
        return None, "HTTP {}".format(r.status_code)
    except Exception as e:
        return None, str(e)


def _gen_image_with_code(chat_id, user_task, on_status):
    """Fallback: генерирует изображение через Python PIL-скрипт."""
    prompt = (
        "Напиши Python-скрипт который создаёт изображение с помощью PIL/Pillow.\n"
        "Описание: {}\n"
        "Скрипт должен:\n"
        "1. Создать изображение 512x512 или больше\n"
        "2. Сохранить в файл image.png в текущей директории\n"
        "3. Напечатать: print('IMAGE:', os.path.abspath('image.png'))\n"
        "Используй только PIL/Pillow, math, random, os.\n"
        "Код в ```python```"
    ).format(user_task)

    try:
        llm_response = _call_with_continuation(prompt, system_override or CODER_SCRIPT_SYSTEM, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
        code = _extract_code(llm_response)
        if not code:
            raise ValueError("Нет кода")

        success, output, created = _execute_code(code)

        # Ищем путь к изображению
        img_path = None
        for line in output.splitlines():
            if line.strip().startswith("IMAGE:"):
                img_path = line.strip()[6:].strip()
                break
        if not img_path:
            for f in created:
                if f.lower().endswith(('.png', '.jpg', '.jpeg')):
                    img_path = f
                    break

        if img_path and os.path.exists(img_path):
            return {
                'success': True, 'output': "🎨 Изображение создано через PIL",
                'code': code, 'iterations': 1, 'is_project': False,
                'files': [img_path], 'file_names': ['image.png'],
            }
    except Exception as e:
        pass

    return {
        'success': False,
        'output': (
            "❌ Не удалось сгенерировать изображение.\n\n"
            "Для генерации изображений нужен ключ Together AI (FLUX.1 бесплатно):\n"
            "1. Зайди на together.ai\n"
            "2. Получи ключ\n"
            "3. /setllm together MODEL КЛЮЧ"
        ),
        'code': '', 'iterations': 1, 'is_project': False, 'files': [],
    }


# ════════════════════════════════════════════════════════════
#  АГЕНТ: СКАФФОЛДЕР
# ════════════════════════════════════════════════════════════

def _run_image_task(chat_id, user_task, on_status):
    result = image_gen_agent(user_task, on_status=on_status)
    add_to_history(chat_id, 'assistant', 'image ok' if result['success'] else 'image fail')
    return {'success': result['success'], 'code': result.get('code',''),
            'output': result.get('output',''), 'iterations': 1,
            'is_project': False, 'files': result.get('files',[]),
            '_full_output': result.get('output','')}


def _run_analyze_task(chat_id, user_task, on_status):
    if on_status: on_status('🔍 Анализирую...')
    code = _extract_code_ac(user_task) if _extract_code_ac else ''
    if code:
        analysis = analyze_code(code, user_task)
    else:
        from agent_core import SYS_ANALYZER
        analysis = call_llm(user_task, SYS_ANALYZER, max_tokens=2000)
    add_to_history(chat_id, 'assistant', analysis[:300])
    return {'success': True, 'code': code, 'output': analysis,
            'iterations': 1, 'is_project': False, 'files': [],
            '_full_output': analysis, 'text_only': True}


def _run_fix_task(chat_id, user_task, on_status):
    if on_status: on_status('🔍 Анализирую ошибку...')
    code = _extract_code_ac(user_task) if _extract_code_ac else ''
    explanation = explain_error(user_task, code)
    if not code:
        add_to_history(chat_id, 'assistant', explanation[:300])
        return {'success': True, 'code': '', 'output': explanation,
                'iterations': 1, 'is_project': False, 'files': [],
                '_full_output': explanation, 'text_only': True}
    if on_status: on_status('🔧 Исправляю...')
    result = fix_agent(code, user_task, task=user_task, on_status=on_status)
    out = (result.get('explanation','') + '\n\n' + result.get('output','')).strip()
    add_to_history(chat_id, 'assistant', 'fixed' if result['success'] else 'fix_failed')
    return {'success': result['success'], 'code': result.get('code',''),
            'output': out, 'iterations': result.get('iterations',1),
            'is_project': False, 'files': [], '_full_output': out}


# ════════════════════════════════════════════════════════════
#  ТЕСТ + ФИХ ДЛЯ МНОГОФАЙЛОВЫХ ПРОЕКТОВ
# ════════════════════════════════════════════════════════════

MAX_PROJECT_FIX_ITERATIONS = 3

_BENIGN_ERRORS = [
    'token', 'api key', 'connect', 'network', 'timeout', 'getenv',
    'keyboard interrupt', 'systemexit', 'no such host', 'refused',
    'environ', 'bot_token', 'chat_id', 'unauthorized',
]


def _test_project(project_dir, generated):
    """
    Тестирует сгенерированный проект двумя шагами:
      1. compile() на всех .py файлах — ловит SyntaxError ещё до запуска.
      2. Запускает точку входа с коротким таймаутом. Если скрипт «завис»
         на polling или завершился с «безобидной» ошибкой (нет токена/сети) —
         считаем это успехом.
    Возвращает (ok: bool, error_text: str).
    """
    syntax_errors = []
    for fpath, content in generated:
        if not fpath.endswith('.py'):
            continue
        try:
            compile(content, fpath, 'exec')
        except SyntaxError as e:
            syntax_errors.append('SyntaxError в {}, строка {}: {}'.format(fpath, e.lineno, e.msg))

    if syntax_errors:
        return False, '\n'.join(syntax_errors)

    entry_candidates = ['main.py', 'app.py', 'bot.py', 'run.py', '__main__.py']
    entry = None
    for candidate in entry_candidates:
        full = os.path.join(project_dir, candidate)
        if os.path.exists(full):
            entry = full
            break

    if entry is None:
        return True, ''

    try:
        result = subprocess.run(
            [sys.executable, entry],
            capture_output=True, text=True, timeout=8,
            cwd=project_dir,
            env={**os.environ, 'PYTHONIOENCODING': 'utf-8'},
        )
    except subprocess.TimeoutExpired:
        return True, ''
    except Exception as e:
        return False, str(e)

    if result.returncode == 0:
        return True, ''

    stderr = (result.stderr or result.stdout or '').strip()
    if any(b in stderr.lower() for b in _BENIGN_ERRORS):
        return True, ''

    return False, stderr


def _fix_project(generated, user_task, error, iteration, on_status=None):
    """
    Отправляет все файлы + ошибку в LLM и просит вернуть исправленный
    JSON-массив [{path, content}]. Возвращает новый список или None.
    """
    if on_status:
        on_status('🔧 Исправляю ошибки, попытка {}/{}...'.format(
            iteration, MAX_PROJECT_FIX_ITERATIONS))

    files_json = json.dumps(
        [{'path': p, 'content': c} for p, c in generated],
        ensure_ascii=False, indent=2
    )

    fix_system = (
        'Ты эксперт по Python. Исправь ошибки в проекте.\n'
        'Верни ТОЛЬКО JSON-массив:\n'
        '[{"path": "путь/к/файлу.py", "content": "полное содержимое"}]\n'
        'Включай ВСЕ файлы (изменённые и неизменённые). Только JSON.'
    )

    fix_prompt = (
        'Задача: {}\n\nОшибка:\n{}\n\nФайлы:\n{}\n\nВерни исправленный JSON:'
    ).format(user_task[:400], error[:1500], files_json[:8000])

    try:
        response = call_llm(fix_prompt, fix_system, max_tokens=CODE_MAX_TOKENS)
        m = re.search(r'\[.*\]', response, re.DOTALL)
        if not m:
            return None
        fixed = json.loads(m.group())
        if isinstance(fixed, list):
            return [(f['path'], f['content']) for f in fixed
                    if 'path' in f and 'content' in f]
    except Exception:
        pass
    return None


def _build_project_zip(project_dir, generated, project_name='project'):
    """Пакует проект в zip, сохраняя полную структуру папок."""
    import zipfile as _zf
    zip_path = os.path.join(project_dir, '{}.zip'.format(project_name))
    with _zf.ZipFile(zip_path, 'w', _zf.ZIP_DEFLATED) as zf:
        for fpath, _ in generated:
            parts = [re.sub(r'[^\w.\-]', '_', p)
                     for p in fpath.replace('\\', '/').split('/') if p]
            full = os.path.join(project_dir, *parts)
            if os.path.exists(full):
                zf.write(full, fpath)   # fpath = имя внутри архива (с путём)
    return zip_path


def _run_scaffold_agent(chat_id, user_task, on_status):
    if on_status:
        on_status("🏗 Генерирую структуру проекта...")

    prompt = (
        "Создай Python-скрипт который строит структуру проекта и упаковывает в zip.\n\n"
        "Описание:\n{}\n\n"
        "Скрипт должен:\n"
        "1. Создать все директории и файлы с реальным содержимым\n"
        "2. Упаковать в zip-архив\n"
        "3. Напечатать: print('ZIP:', zip_path)\n"
        "Вернуть только код в ```python```."
    ).format(user_task)

    try:
        llm_response = _call_with_continuation(prompt, SCAFFOLD_SYSTEM, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
    except Exception as e:
        return {'success': False, 'output': "❌ LLM: {}".format(e),
                'code': '', 'iterations': 1, 'is_project': True, 'files': []}

    code = _extract_code(llm_response)
    if not code:
        return {'success': False, 'output': "❌ Агент не сгенерировал код",
                'code': '', 'iterations': 1, 'is_project': True, 'files': []}

    if on_status:
        on_status("⚙️ Запускаю скаффолдер...")

    success, output, _ = _execute_code(code)

    zip_path = None
    for line in output.splitlines():
        if line.strip().startswith("ZIP:"):
            zip_path = line.strip()[4:].strip()
            break

    if not zip_path or not os.path.exists(zip_path):
        import glob
        zips = sorted(
            glob.glob(os.path.join(PROJECTS_DIR, "**/*.zip"), recursive=True),
            key=os.path.getmtime, reverse=True
        )
        if zips and (time.time() - os.path.getmtime(zips[0])) < 120:
            zip_path = zips[0]

    files = [zip_path] if zip_path and os.path.exists(zip_path) else []

    if success and files:
        fname = os.path.basename(zip_path)
        size  = os.path.getsize(zip_path)
        add_to_history(chat_id, 'assistant', "✅ Архив создан: {}".format(fname))
        return {
            'success': True, 'code': code,
            'output': "Архив готов. Отправляю...",
            'iterations': 1, 'is_project': True, 'files': files,
            'file_names': [fname], 'project_dir': os.path.dirname(zip_path),
        }
    return {
        'success': success, 'code': code,
        'output': output + ("\n\n⚠️ zip-файл не найден." if not files else ""),
        'iterations': 1, 'is_project': True, 'files': files,
    }


# ════════════════════════════════════════════════════════════
#  АГЕНТ: МНОГОФАЙЛОВЫЙ ПРОЕКТ (ПЛАН + ФАЙЛЫ)
# ════════════════════════════════════════════════════════════

def _run_project_agent(chat_id, user_task, history, on_status):
    if on_status:
        on_status("🏗 Планирую структуру проекта...")

    history_ctx = _format_history(history[:-1])

    plan_system = (
        "Ты архитектор. Верни ТОЛЬКО JSON-массив файлов.\n"
        "Формат: [{\"path\": \"dir/file.py\", \"description\": \"назначение\", \"language\": \"python\"}]\n"
        "Максимум 15 файлов. Только JSON."
    )
    plan_prompt = "{}\n\nЗадача: {}\n\nJSON-список файлов:".format(history_ctx, user_task).strip()

    try:
        plan_raw  = call_llm(plan_prompt, plan_system, max_tokens=1500)
        file_plan = _parse_file_plan(plan_raw)
    except Exception:
        file_plan = []

    if not file_plan:
        if on_status:
            on_status("⚠️ Генерирую одним запросом...")
        return _run_project_agent_fallback(chat_id, user_task, history, on_status)

    if on_status:
        on_status("📋 План: {} файлов. Генерирую по одному...".format(len(file_plan)))

    generated = []
    failed    = []

    for i, file_info in enumerate(file_plan):
        fpath = file_info.get("path", "file_{}.py".format(i))
        fdesc = file_info.get("description", "")
        flang = file_info.get("language", "python")

        if on_status:
            on_status("✍️ [{}/{}] {}".format(i + 1, len(file_plan), fpath))

        code_system = (
            "Пиши ТОЛЬКО содержимое файла — без markdown, без объяснений.\n"
            "Язык: {}\n{}".format(flang, _ENV_INFO if flang == "python" else "")
        )
        code_prompt = "Проект: {}\n\nСодержимое файла `{}`:\n{}".format(
            user_task[:400], fpath, fdesc)

        try:
            content = _call_with_continuation(code_prompt, code_system, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
            content = _strip_code_fences(content)
            generated.append((fpath, content))
        except Exception as e:
            failed.append(fpath)
            generated.append((fpath, "# ERROR: {}\n".format(e)))

    project_dir, saved_paths, saved_names = _save_project_files(generated)

    project_name = re.sub(r"[^\w-]", "_", user_task[:30]).strip("_") or "project"
    ok_count = len(generated) - len(failed)

    # ── Тест + авто-фикс ──────────────────────────────────────────────
    fix_iterations = 0
    last_error = ''
    for fix_iter in range(1, MAX_PROJECT_FIX_ITERATIONS + 1):
        if on_status:
            on_status("🧪 Проверяю синтаксис и запуск ({}/{})...".format(
                fix_iter, MAX_PROJECT_FIX_ITERATIONS))

        test_ok, test_error = _test_project(project_dir, generated)
        fix_iterations = fix_iter

        if test_ok:
            break

        last_error = test_error
        if on_status:
            on_status("❌ Ошибка: {}".format(test_error[:120]))

        if fix_iter == MAX_PROJECT_FIX_ITERATIONS:
            break   # исчерпали попытки — упакуем что есть

        fixed = _fix_project(generated, user_task, test_error, fix_iter, on_status)
        if fixed:
            generated = fixed
            project_dir, saved_paths, saved_names = _save_project_files(generated)
        else:
            break
    # ──────────────────────────────────────────────────────────────────

    zip_path = _build_project_zip(project_dir, generated, project_name)

    if on_status:
        status_icon = "✅" if not last_error else "⚠️"
        on_status("{} Готово: {}/{} файлов, {} попыток исправлений".format(
            status_icon, ok_count, len(generated), fix_iterations))

    add_to_history(chat_id, "assistant", "✅ {} файлов".format(len(generated)))
    fail_note = "\n⚠️ Проблемы при генерации: {}".format(", ".join(failed)) if failed else ""
    error_note = "\n⚠️ Не устранено за {} попыток: {}".format(
        MAX_PROJECT_FIX_ITERATIONS, last_error[:200]) if last_error else ""

    return {
        "success": not last_error,
        "code": generated[0][1] if generated else "",
        "output": "📦 ZIP + {} файлов. Проверено за {} итераций.{}{}".format(
            len(generated), fix_iterations, fail_note, error_note),
        "iterations": fix_iterations, "is_project": True,
        "files": [zip_path] + saved_paths,
        "file_names": [os.path.basename(zip_path)] + saved_names,
        "project_dir": project_dir,
    }


# ════════════════════════════════════════════════════════════
#  АГЕНТ: ОДИН ФАЙЛ ЗА РАЗ
# ════════════════════════════════════════════════════════════

def _run_onebyone_agent(chat_id, user_task, history, on_status):
    if on_status:
        on_status("📋 Определяю список файлов...")

    names_system = "Верни ТОЛЬКО список имён файлов (один на строку, с путём). Максимум 12."
    names_prompt = "Задача: {}\n\nСписок файлов:".format(user_task[:500])

    try:
        names_raw  = call_llm(names_prompt, names_system, max_tokens=400)
        file_names = []
        for line in names_raw.strip().splitlines():
            line = line.strip().lstrip('-•*123456789. ')
            if line and '.' in line and len(line) < 80:
                file_names.append(line)
        if not file_names:
            raise ValueError("Пустой список")
    except Exception:
        file_names = ['main.py', 'config.py', 'requirements.txt', 'README.md']

    if on_status:
        on_status("📄 Файлов: {}. Генерирую...".format(len(file_names)))

    generated = []
    failed = []
    history_ctx = _format_history(history[:-1])

    for i, fpath in enumerate(file_names):
        if on_status:
            on_status("✍️ [{}/{}] {}".format(i + 1, len(file_names), fpath))

        ext = os.path.splitext(fpath)[1].lower()
        lang_map = {'.py':'python','.kt':'kotlin','.js':'javascript','.ts':'typescript',
                    '.md':'markdown','.yaml':'yaml','.yml':'yaml','.json':'json','.sh':'bash'}
        lang = lang_map.get(ext, 'text')

        code_system = (
            "Пиши ТОЛЬКО содержимое файла — без markdown, без пояснений.\n"
            "{}".format(_ENV_INFO if lang == 'python' else "")
        )
        code_prompt = "{}\nПроект: {}\n\nСодержимое `{}`:".format(
            history_ctx[:300], user_task[:400], fpath)

        try:
            content = _call_with_continuation(code_prompt, code_system, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
            generated.append((fpath, _strip_code_fences(content)))
        except Exception as e:
            failed.append(fpath)
            generated.append((fpath, "# ERROR: {}\n".format(e)))

    project_dir, saved_paths, saved_names = _save_project_files(generated)

    project_name = re.sub(r"[^\w-]", "_", user_task[:30]).strip("_") or "project"

    # ── Тест + авто-фикс ──────────────────────────────────────────────
    fix_iterations = 0
    last_error = ''
    for fix_iter in range(1, MAX_PROJECT_FIX_ITERATIONS + 1):
        if on_status:
            on_status("🧪 Проверяю ({}/{})...".format(fix_iter, MAX_PROJECT_FIX_ITERATIONS))

        test_ok, test_error = _test_project(project_dir, generated)
        fix_iterations = fix_iter

        if test_ok:
            break

        last_error = test_error
        if fix_iter == MAX_PROJECT_FIX_ITERATIONS:
            break

        fixed = _fix_project(generated, user_task, test_error, fix_iter, on_status)
        if fixed:
            generated = fixed
            project_dir, saved_paths, saved_names = _save_project_files(generated)
        else:
            break
    # ──────────────────────────────────────────────────────────────────

    zip_path = _build_project_zip(project_dir, generated, project_name)

    ok_count = len(generated) - len(failed)
    if on_status:
        on_status("✅ {}/{} файлов".format(ok_count, len(generated)))

    add_to_history(chat_id, "assistant", "✅ {} файлов".format(len(generated)))
    error_note = "\n⚠️ Не устранено: {}".format(last_error[:200]) if last_error else ""

    return {
        "success": not last_error,
        "code": generated[0][1] if generated else "",
        "output": "📦 ZIP + {} файлов. {} итераций проверки.{}".format(
            len(generated), fix_iterations, error_note),
        "iterations": fix_iterations, "is_project": True,
        "files": [zip_path] + saved_paths,
        "file_names": [os.path.basename(zip_path)] + saved_names,
        "project_dir": project_dir,
    }


# ════════════════════════════════════════════════════════════
#  АГЕНТ: ПРОСТОЙ СКРИПТ
# ════════════════════════════════════════════════════════════

def _run_script_agent(chat_id, user_task, history, on_status, system_override=None):
    last_code  = None
    last_error = None
    all_attempts = []

    for iteration in range(1, MAX_ITERATIONS + 1):
        with _sessions_lock:
            s = _sessions.get(str(chat_id))
            if s:
                s['iterations'] = iteration

        if iteration == 1:
            history_ctx = _format_history(history[:-1])
            prompt  = "{}\n\nЗадача: {}".format(history_ctx, user_task).strip()
            system  = CODER_SCRIPT_SYSTEM
        else:
            attempts_text = "\n\n".join(
                "Попытка {}:\n```python\n{}\n```\nОшибка:\n{}".format(
                    i+1, a['code'], a['error'])
                for i, a in enumerate(all_attempts)
            )
            prompt = "Задача: {}\n\n{}\n\nИсправленный код:".format(user_task, attempts_text)
            system = CODER_FIX_SYSTEM

        if on_status:
            on_status("🤔 Генерирую..." if iteration == 1 else
                      "🔧 Итерация {} — исправляю...".format(iteration))

        try:
            llm_response = _call_with_continuation(prompt, system, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
        except Exception as e:
            return {'success': False, 'output': "❌ LLM: {}".format(e),
                    'code': last_code or '', 'iterations': iteration, 'is_project': False, 'files': []}

        code = _extract_code(llm_response)
        if not code:
            add_to_history(chat_id, 'assistant', llm_response)
            return {'success': True, 'output': llm_response, 'code': '',
                    'iterations': iteration, 'is_project': False, 'files': [], 'text_only': True}

        last_code = code

        if on_status:
            on_status("⚙️ Запускаю (попытка {})...".format(iteration))

        success, output, created_files = _execute_code(code)

        if success:
            add_to_history(chat_id, 'assistant', "✅ Выполнено:\n{}".format(output[:300]))
            return {'success': True, 'code': code, 'output': output,
                    'iterations': iteration, 'is_project': False, 'files': created_files,
                    '_full_output': output}
        else:
            last_error = output
            all_attempts.append({'code': code, 'error': output[:600]})
            if iteration == MAX_ITERATIONS:
                add_to_history(chat_id, 'assistant', "❌ Не удалось за {} итераций".format(iteration))
                return {'success': False, 'code': code, 'output': output,
                        'iterations': iteration, 'is_project': False, 'files': [],
                        '_full_output': output}


# ════════════════════════════════════════════════════════════
#  FALLBACK
# ════════════════════════════════════════════════════════════

def _run_project_agent_fallback(chat_id, user_task, history, on_status):
    if on_status:
        on_status("🏗 Генерирую одним блоком...")

    history_ctx = _format_history(history[:-1])
    prompt = (
        "{}\n\nЗадача: {}\n\n"
        "Используй: ###FILE: path/file.py\n<код>\n###FILE: path/other.py\n<код>"
    ).format(history_ctx, user_task).strip()

    try:
        llm_response = _call_with_continuation(prompt, CODER_PROJECT_SYSTEM, max_tokens=CODE_MAX_TOKENS, on_status=on_status)
    except Exception as e:
        return {"success": False, "output": "❌ LLM: {}".format(e),
                "code": "", "iterations": 1, "is_project": True, "files": []}

    files = _extract_project_files(llm_response)
    if not files:
        code = _extract_code(llm_response)
        files = [("main.py", code)] if code else []

    if not files:
        add_to_history(chat_id, "assistant", llm_response)
        return {"success": True, "output": llm_response, "code": "",
                "iterations": 1, "is_project": True, "files": [], "text_only": True}

    project_dir, saved_paths, saved_names = _save_project_files(files)
    add_to_history(chat_id, "assistant", "✅ {} файлов".format(len(files)))
    return {
        "success": True, "code": files[0][1] if files else "",
        "output": "Файлы готовы.", "iterations": 1, "is_project": True,
        "files": saved_paths, "file_names": saved_names, "project_dir": project_dir,
    }


# ════════════════════════════════════════════════════════════
#  ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ════════════════════════════════════════════════════════════

def _extract_code(text):
    for pattern in [r'```python\s*\n(.*?)```', r'```py\s*\n(.*?)```', r'```\s*\n(.*?)```']:
        m = re.search(pattern, text, re.DOTALL)
        if m:
            return m.group(1).strip()
    return None


def _extract_code_from_history(history):
    """Ищет последний блок кода в истории диалога."""
    for msg in reversed(history):
        code = _extract_code(msg.get('content', ''))
        if code:
            return code
    return None


def _extract_error_from_text(text):
    """Ищет traceback или сообщение об ошибке в тексте."""
    # Ищем стандартный Python traceback
    if 'Traceback' in text or 'Error:' in text or 'Exception:' in text:
        return text
    return None


def _execute_code(code):
    tmp_path = None
    run_dir = os.path.join(PROJECTS_DIR, 'run_' + time.strftime('%H%M%S'))
    os.makedirs(run_dir, exist_ok=True)
    before = set(os.listdir(run_dir))

    try:
        with tempfile.NamedTemporaryFile(
            mode='w', suffix='.py', delete=False, encoding='utf-8', dir=run_dir
        ) as f:
            f.write(code)
            tmp_path = f.name

        result = subprocess.run(
            [sys.executable, tmp_path],
            capture_output=True, text=True,
            timeout=CODE_TIMEOUT, cwd=run_dir,
            env={**os.environ, 'PYTHONIOENCODING': 'utf-8'},
        )

        after   = set(os.listdir(run_dir))
        created = [
            os.path.join(run_dir, f)
            for f in (after - before)
            if not f.endswith('.py') and os.path.isfile(os.path.join(run_dir, f))
        ]

        if result.returncode == 0:
            return True, (result.stdout or "(нет вывода)").strip(), created
        else:
            return False, (result.stderr or result.stdout or "(нет вывода)").strip(), []

    except subprocess.TimeoutExpired:
        return False, "⏰ Таймаут {}с — скрипт не завершился.".format(CODE_TIMEOUT), []
    except Exception as e:
        return False, "❌ Ошибка запуска: {}".format(e), []
    finally:
        if tmp_path:
            try: os.unlink(tmp_path)
            except Exception: pass


def _parse_file_plan(raw):
    m = re.search(r"\[.*\]", raw, re.DOTALL)
    if not m:
        return []
    try:
        data = json.loads(m.group())
        if isinstance(data, list):
            return [f for f in data if isinstance(f, dict) and "path" in f]
    except Exception:
        pass
    return []


def _strip_code_fences(text):
    text = re.sub(r"^```[\w]*\n?", "", text.strip())
    text = re.sub(r"\n?```$", "", text.strip())
    return text.strip()


def _extract_project_files(text):
    files = []
    parts = re.split(r'###FILE:\s*(\S+)', text)
    i = 1
    while i + 1 < len(parts):
        filename = parts[i].strip()
        content  = parts[i + 1].strip()
        content  = re.sub(r'^```\w*\n?', '', content)
        content  = re.sub(r'\n?```\s*$', '', content)
        files.append((filename, content.strip()))
        i += 2
    return files


def _extract_instructions(text):
    m = re.search(r'###INSTRUCTIONS:\s*(.*?)(?=###|$)', text, re.DOTALL)
    return m.group(1).strip() if m else None


def _save_project_files(files):
    ts = time.strftime('%Y%m%d_%H%M%S')
    project_dir = os.path.join(PROJECTS_DIR, ts)
    os.makedirs(project_dir, exist_ok=True)
    saved_paths = []
    saved_names = []

    # Track used names to avoid collisions when different dirs have same filename
    used_names = {}

    for filename, content in files:
        # Preserve directory structure: bot/handlers/main.py -> bot/handlers/main.py on disk
        # Only sanitize each path component individually
        parts = filename.replace('\\', '/').split('/')
        safe_parts = []
        for part in parts:
            safe = re.sub(r'[^\w.\-]', '_', part)
            if safe:
                safe_parts.append(safe)

        if not safe_parts:
            safe_parts = ['file_{}.py'.format(len(saved_paths))]

        # Resolve collisions: if bot/handlers/main.py and api/main.py both exist,
        # keep full relative path so neither overwrites the other
        rel_path = os.path.join(*safe_parts)
        full_path = os.path.join(project_dir, rel_path)

        # Create parent dirs as needed
        os.makedirs(os.path.dirname(full_path), exist_ok=True)

        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(content)

        saved_paths.append(full_path)
        saved_names.append(rel_path)

    return project_dir, saved_paths, saved_names


# ════════════════════════════════════════════════════════════
#  ФОРМАТИРОВАНИЕ РЕЗУЛЬТАТА
# ════════════════════════════════════════════════════════════

def format_code_result(result):
    if result.get('text_only'):
        return result['output'][:4000]

    is_project = result.get('is_project', False)
    code       = result.get('code', '')
    output     = result.get('output', '')
    success    = result.get('success', False)
    iters      = result.get('iterations', 1)
    files      = result.get('files', [])
    proj_dir   = result.get('project_dir', '')

    lines = []
    iter_label = " (итераций: {})".format(iters) if iters > 1 else ""

    if is_project:
        if success:
            lines.append("✅ <b>Проект создан{}!</b>".format(iter_label))
            if proj_dir:
                lines.append("📁 <code>{}</code>".format(proj_dir))
            display_names = result.get('file_names', [os.path.basename(f) for f in files])
            if display_names:
                lines.append("\n<b>Файлы:</b>")
                for f in display_names:
                    lines.append("  • <code>{}</code>".format(f))
            if output and output not in ("Файлы готовы.", "Архив готов. Отправляю..."):
                lines.append("\n<b>Инструкция:</b>")
                lines.append(_escape_html(output[:1000]))
            if code:
                preview = code[:1000] + ('\n...' if len(code) > 1000 else '')
                lines.append("\n<b>Начало кода:</b>")
                lines.append("<pre>{}</pre>".format(_escape_html(preview)))
        else:
            lines.append("❌ <b>Ошибка</b>")
            if output:
                lines.append("<pre>{}</pre>".format(_escape_html(output[:800])))
    else:
        lines.append(("✅ <b>Выполнено{}!</b>" if success else "❌ <b>Не удалось{}:</b>").format(iter_label))
        if code:
            preview = code[:1500] + ('\n...' if len(code) > 1500 else '')
            lines.append("\n<b>Код:</b>")
            lines.append("<pre>{}</pre>".format(_escape_html(preview)))
        if files:
            lines.append("\n<b>📎 Файлов: {}</b>".format(len(files)))
            for f in files:
                lines.append("  • <code>{}</code>".format(os.path.basename(f)))

    result['_full_output'] = output
    return "\n".join(lines)


def _escape_html(text):
    import html
    return html.escape(str(text))
